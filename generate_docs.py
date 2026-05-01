"""Generate tentative-implementation-plan.docx and tentative-implementation-plan.xlsx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ============================================================
# WORD DOCUMENT
# ============================================================

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table_row(table, cells_data, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    return row

# Title Page
title = doc.add_heading('SNM Portal', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Tentative Implementation Plan', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Date: March 28, 2026').bold = True
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Status: Draft / Tentative')

doc.add_page_break()

# 1. Context
add_heading('1. Context', level=1)
doc.add_paragraph(
    'Build a multi-tenant B2B portal for managing customers, enquiries, quotations, and costing '
    '- with role-based access, dynamic menus, and email capabilities. The system handles an '
    'Enquiry > Costing > Quotation workflow for steel/metal products.'
)

# Tech Stack
add_heading('1.1 Technology Stack', level=2)
tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = 'Light Grid Accent 1'
tech_table.rows[0].cells[0].text = 'Component'
tech_table.rows[0].cells[1].text = 'Technology'
for r in tech_table.rows[0].cells:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

tech_items = [
    ('Backend', 'FastAPI (Python)'),
    ('Frontend', 'Angular 19 (Standalone Components)'),
    ('UI Library', 'Angular Material'),
    ('Database', 'SQL Server (Azure SQL in production)'),
    ('ORM', 'SQLAlchemy'),
    ('Migrations', 'Alembic'),
    ('Authentication', 'JWT (Access + Refresh Tokens)'),
    ('File Storage', 'Azure Blob Storage'),
    ('Email', 'SMTP (per-company, optional/configurable)'),
]
for comp, tech in tech_items:
    add_table_row(tech_table, [comp, tech])

doc.add_paragraph('')

# 2. Schema Changes
add_heading('2. Schema Changes (vs. Original)', level=1)

add_heading('2.1 New Table: UserRoleMap', level=2)
doc.add_paragraph(
    'A user can be mapped to multiple companies, each with a different role. '
    'On login, the user gets a list of their companies. The isDefault flag auto-selects the primary company. '
    'A company switcher in the header allows switching at runtime.'
)
urm_table = doc.add_table(rows=1, cols=3)
urm_table.style = 'Light Grid Accent 1'
urm_table.rows[0].cells[0].text = 'Field'
urm_table.rows[0].cells[1].text = 'Type'
urm_table.rows[0].cells[2].text = 'Description'
for c in urm_table.rows[0].cells:
    for p in c.paragraphs:
        for run in p.runs:
            run.bold = True
urm_fields = [
    ('userRoleMapId', 'int', 'Auto PK'),
    ('userId', 'int', 'FK -> User'),
    ('roleId', 'int', 'FK -> Role'),
    ('companyId', 'int', 'FK -> Company'),
    ('isDefault', 'bit', 'Default company on login'),
    ('createdon', 'datetime', ''),
    ('createdby', 'int', ''),
    ('isActive', 'bit', ''),
]
for f in urm_fields:
    add_table_row(urm_table, f)

add_heading('2.2 Version Columns Added', level=2)
add_bullet(' - add versionNo int DEFAULT 1', bold_prefix='CustomerEnquiryCosting')
add_bullet(' - add versionNo int DEFAULT 1, parentQuotId int NULL (self-FK)', bold_prefix='QuotSummary')

doc.add_paragraph(
    'Versioning: When a new version is created, a new row is inserted with versionNo + 1. '
    'For quotations, parentQuotId points to the original quotId. Latest version = MAX(versionNo) per parent. '
    'Previous versions become read-only.'
)

add_heading('2.3 Menu Tree Clarification', level=2)
add_bullet('MenuMaster.parentMenuId supports unlimited nesting (self-referential FK)')
add_bullet('A "menu" represents any resource/component (not just navigation items)')
add_bullet('API returns menus as a recursive tree structure')
add_bullet('Role-Menu Mapping UI shows full tree with expand/collapse and CRUD checkboxes at each node')

add_heading('2.4 Other Schema Fixes', level=2)
add_bullet('itemSize table corrected: itemSizeId (int PK), itemId (int FK), itemSize (varchar 50)')
add_bullet('userPassword column: increased to varchar(255) for bcrypt hash storage')

doc.add_page_break()

# 3. Project Structure
add_heading('3. Project Structure', level=1)

add_heading('3.1 Backend Structure', level=2)
backend_items = [
    'app/core/ - config.py, security.py, dependencies.py, email.py',
    'app/models/ - SQLAlchemy ORM models (25 tables)',
    'app/schemas/ - Pydantic request/response schemas',
    'app/api/v1/ - Route handlers (auth, company, users, roles, menus, masters, customers, enquiries, quotations, assets, email)',
    'app/services/ - Business logic (auth, user, menu, enquiry, quotation, costing, versioning, email, azure_blob)',
    'alembic/ - Database migration scripts',
    '.env / .env.example - Environment configuration',
]
for item in backend_items:
    add_bullet(item)

add_heading('3.2 Frontend Structure', level=2)
frontend_items = [
    'core/auth/ - auth.service, auth.guard, auth.interceptor, token.service',
    'core/services/ - api.service, menu.service, company-context.service, notification.service',
    'shared/components/ - dynamic-menu, data-table, confirm-dialog, company-switcher, page-header',
    'shared/directives/ - has-permission.directive (*hasPermission)',
    'layout/ - main-layout (sidenav + toolbar), header (with company switcher), sidebar',
    'features/auth/ - Login page with company picker',
    'features/dashboard/ - Dashboard',
    'features/company/ - Company management (Super Admin)',
    'features/users/ - User CRUD + company-role assignment',
    'features/roles/ - Role CRUD + role-menu-mapping (tree view)',
    'features/masters/ - 11 master modules (item-grade, item-name, delivery-term, etc.)',
    'features/customers/ - Customer list, form (tabbed: Info, Contacts, Sites)',
    'features/enquiries/ - Enquiry list, form, details, costing (versioned)',
    'features/quotations/ - Quotation list, form, details, T&C, version-history, print/PDF',
    'features/assets/ - Asset upload/download',
]
for item in frontend_items:
    add_bullet(item)

doc.add_page_break()

# 4. Implementation Phases
add_heading('4. Implementation Phases', level=1)

phases = [
    ('Phase 1: Project Setup & Foundation', [
        'Backend scaffolding - FastAPI project, folder structure, .env config',
        'DB connection - SQLAlchemy async engine with pyodbc for SQL Server (connection string from .env)',
        'Alembic setup - Configure for SQL Server, initial migration',
        'Frontend scaffolding - ng new with Angular 19 standalone components, Angular Material, routing',
        'CORS configuration - Allow Angular dev server',
    ]),
    ('Phase 2: Core Models & Auth', [
        'SQLAlchemy models - All 25 tables as ORM models with relationships',
        'Alembic migration - Generate initial migration from models',
        'Auth system - JWT (access + refresh tokens), password hashing (bcrypt), login endpoint',
        'Login flow: credentials -> validate -> return companies list -> company picker (if multiple) -> issue JWT with { user_id, company_id, role_id, is_super_admin }',
        'Company switching - POST /api/v1/auth/switch-company re-issues JWT',
        'Auth middleware - get_current_user dependency, company-scoped queries',
        'Frontend auth - Login page, company picker, token storage, HTTP interceptor, auth guard',
    ]),
    ('Phase 3: Company, User & Role Management', [
        'Company CRUD - API + Angular form (Super Admin only)',
        'Role CRUD - API + Angular form (company-scoped)',
        'Menu Master CRUD - Parent-child tree management, drag-drop reorder',
        'Menu Tree API - GET /api/v1/menus/tree returns recursive JSON',
        'Role-Menu Mapping - mat-tree UI with 4 CRUD checkboxes per node, parent toggles children',
        'User CRUD - API + Angular form',
        'User-Role-Company Mapping - Assign user to one or more companies with role per company',
        'Dynamic sidebar - Menu tree filtered by user role permissions (CanRead=true)',
        'Company Switcher - Dropdown in header, switching reloads menu & data',
        'Permission directive - *hasPermission to show/hide UI elements',
    ]),
    ('Phase 4: Master Data Modules', [
        'Item Grade - CRUD (list + add/edit dialog + soft delete)',
        'Item Name - CRUD (linked to grade)',
        'Item Length - CRUD (linked to item)',
        'Item Size - CRUD (linked to item)',
        'Delivery Term - CRUD',
        'Delivery Mode - CRUD',
        'Contact Type - CRUD',
        'Customer Classification - CRUD',
        'Cost Point Master - CRUD (isPrimary, isTax flags)',
        'Terms & Conditions Master - CRUD',
        'Raw Material Cost - CRUD (dia, tpcost, effectedFrom)',
    ]),
    ('Phase 5: Customer Management', [
        'Customer Master - List with search/filter, Add/Edit form',
        'Customer Contacts - Sub-table under customer',
        'Customer Sites - Sub-table under customer (delivery addresses)',
        'Tabbed UI - Customer form with tabs: Basic Info | Contacts | Sites',
    ]),
    ('Phase 6: Enquiry Management', [
        'Enquiry List - Filterable list with status indicators',
        'Enquiry Form - Select customer, auto-populate contacts/sites',
        'Enquiry Details - Add line items as editable table rows',
        'Enquiry Costing - Per line-item cost breakdown (20+ cost heads), auto-calculate totals',
        'Costing auto-fill - Pull TP cost from RawMaterialCost based on dia',
        'Costing Versioning - Create New Version button, version history dropdown, read-only old versions',
    ]),
    ('Phase 7: Quotation Management', [
        'Quotation List - Filterable, sortable (shows latest version)',
        'Quotation Form - Link to enquiry, customer, site, contact; delivery term/mode',
        'Quotation Details - Line items with rates, GST, totals',
        'Quotation T&C - Pull from master, allow per-quotation edit',
        'Quotation Versioning - Revise button, QUOT-001-R1 format, version history panel, read-only old versions',
        'Quotation Print/PDF - Formatted printable view (specific version)',
        'Approval workflow - approvedby/approvedon, only latest version can be approved',
    ]),
    ('Phase 8: Asset Management & Email', [
        'Azure Blob Service - Upload/download/delete files',
        'Asset Upload - Link files to enquiry or quotation (specific version)',
        'Asset List - View/download assets per enquiry/quotation',
        'Email Configuration (Super Admin Panel) - SMTP settings optional per company, From=company MailFrom, To=selected contact email',
        'Mail Template Management - Super Admin creates/edits HTML templates per company',
        'Send Email - Optional action on quotation, sends PDF via company SMTP',
    ]),
]

for phase_name, tasks in phases:
    add_heading(phase_name, level=2)
    for i, task in enumerate(tasks, 1):
        doc.add_paragraph(f'{i}. {task}')

doc.add_page_break()

# 5. Key Architecture Decisions
add_heading('5. Key Architecture Decisions', level=1)

add_heading('5.1 Backend', level=2)
backend_decisions = [
    'Multi-tenancy: Every query filtered by companyId from JWT token\'s active company',
    'User <-> Company: Many-to-many via UserRoleMap; user gets different role per company',
    'Company context: Active companyId stored in JWT; switch-company re-issues token',
    'Soft delete: All records use isActive flag, never hard-deleted',
    'Audit fields: createdon, createdby, lastupdateon, lastupdateby auto-populated via SQLAlchemy event listeners',
    'Password storage: bcrypt hashing (varchar 255)',
    'JWT payload: { user_id, company_id, role_id, is_super_admin }',
    'API versioning: /api/v1/ prefix',
    'Service layer: Business logic separated from route handlers',
    'Versioning: Same-table approach with versionNo column; latest = max version per parent entity',
    'DB config from .env: Local dev uses local/dev instance, production uses Azure SQL Server',
    'Email is optional: SMTP settings per company are nullable; features disabled when not configured',
    'Email flow: From = company MailFrom, To = selected customer contact email',
]
for d in backend_decisions:
    add_bullet(d)

add_heading('5.2 Frontend', level=2)
frontend_decisions = [
    'Angular 19 standalone components (no NgModules)',
    'Lazy-loaded feature routes for each module',
    'Angular Material for UI (sidenav, mat-tree, tables, dialogs, forms)',
    'Reactive Forms for all forms',
    'HTTP Interceptor for JWT token injection + 401 redirect',
    'Company Switcher in header toolbar - switching reloads sidebar menu and all data',
    'Permission directive reads RoleMenuMap to conditionally render buttons',
    'Dynamic menu - recursive mat-tree sidebar from server menu tree response',
    'Role-Menu Mapping - mat-tree with checkbox nodes for CRUD permissions per menu',
]
for d in frontend_decisions:
    add_bullet(d)

add_heading('5.3 Database', level=2)
db_decisions = [
    'SQL Server via pyodbc driver (mssql+pyodbc connection string from .env)',
    'Dev: local SQL Server or Azure SQL dev instance; Production: Azure SQL Server',
    'All IDs auto-increment (Identity columns)',
    'Foreign keys as documented in schema',
    'itemSize table corrected, userPassword increased to varchar(255)',
    'New table: UserRoleMap for multi-company user mapping',
    'New columns: versionNo on QuotSummary & CustomerEnquiryCosting; parentQuotId on QuotSummary',
]
for d in db_decisions:
    add_bullet(d)

doc.add_page_break()

# 6. Verification Plan
add_heading('6. Verification Plan', level=1)
verification = [
    'Backend: Run pytest for each API endpoint (auth, CRUD, permissions, versioning)',
    'Frontend: ng serve - verify login flow, company switching, menu rendering, CRUD operations',
    'DB: Run Alembic migrations on SQL Server, verify all 25 tables created',
    'Auth flow: Login -> company picker (multi-company user) -> JWT issued -> company switch -> menu reloads',
    'Permission flow: Role-menu tree assignment -> sidebar reflects only permitted menus -> CRUD buttons hidden per permission',
    'Versioning flow: Create quotation -> revise -> verify version history -> view old version (read-only) -> print specific version',
    'E2E flow: Login -> Create Customer -> Create Enquiry -> Add Costing -> Generate Quotation -> Revise Quotation -> Upload Asset -> Send Email',
]
for i, v in enumerate(verification, 1):
    doc.add_paragraph(f'{i}. {v}')

doc.save('tentative-implementation-plan.docx')
print("Created: tentative-implementation-plan.docx")

# ============================================================
# EXCEL DOCUMENT
# ============================================================

wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Implementation Tasks"

# Styles
header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
header_fill = PatternFill(start_color='1A478A', end_color='1A478A', fill_type='solid')
phase_font = Font(name='Calibri', size=11, bold=True, color='1A478A')
phase_fill = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
normal_font = Font(name='Calibri', size=10)
wrap_alignment = Alignment(wrap_text=True, vertical='top')
center_alignment = Alignment(horizontal='center', vertical='top', wrap_text=True)
thin_border = Border(
    left=Side(style='thin'),
    right=Side(style='thin'),
    top=Side(style='thin'),
    bottom=Side(style='thin')
)

# Headers
headers = ['#', 'Phase', 'Task', 'Sub-Tasks / Details', 'Layer', 'Key Files', 'Prompt Hint (for AI)', 'Status']
col_widths = [5, 20, 35, 55, 10, 35, 55, 12]

for col_num, (header, width) in enumerate(zip(headers, col_widths), 1):
    cell = ws.cell(row=1, column=col_num, value=header)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = center_alignment
    cell.border = thin_border
    ws.column_dimensions[get_column_letter(col_num)].width = width

# Freeze top row
ws.freeze_panes = 'A2'

# Task data
tasks = [
    # Phase 1
    (1, 'Phase 1: Setup', 'Backend scaffolding',
     'Create FastAPI project structure, folder hierarchy, requirements.txt, .env.example with DB_CONNECTION_STRING, JWT_SECRET, AZURE_BLOB_CONNECTION_STRING, AZURE_BLOB_CONTAINER',
     'Backend', 'backend/app/main.py, backend/requirements.txt, backend/.env.example',
     'Create a FastAPI backend project with folder structure: app/core/, app/models/, app/schemas/, app/api/v1/, app/services/. Include .env.example with all config vars. Use pydantic-settings for config.',
     'To Do'),
    (2, 'Phase 1: Setup', 'DB connection setup',
     'SQLAlchemy engine config for SQL Server using pyodbc. Connection string from .env. Async session maker. Dev: local SQL Server. Prod: Azure SQL Server.',
     'Backend', 'backend/app/core/config.py, backend/app/core/dependencies.py',
     'Set up SQLAlchemy with SQL Server using mssql+pyodbc. Load connection string from .env. Create async engine, sessionmaker, and get_db dependency. Support both local and Azure SQL Server.',
     'To Do'),
    (3, 'Phase 1: Setup', 'Alembic setup',
     'Initialize Alembic for SQL Server. Configure env.py to read DB URL from .env.',
     'Backend', 'backend/alembic/, backend/alembic.ini, backend/alembic/env.py',
     'Initialize Alembic in the backend project. Configure alembic.ini and env.py to use SQL Server connection string from .env. Set up for auto-generating migrations from SQLAlchemy models.',
     'To Do'),
    (4, 'Phase 1: Setup', 'Frontend scaffolding',
     'ng new with Angular 19 standalone components. Install Angular Material. Set up routing, environments, proxy config for API.',
     'Frontend', 'frontend/src/app/app.config.ts, frontend/src/app/app.routes.ts, frontend/angular.json',
     'Create Angular 19 project with standalone components (no NgModules). Install Angular Material. Set up app.routes.ts with lazy-loaded routes. Configure proxy.conf.json to forward /api to FastAPI backend.',
     'To Do'),
    (5, 'Phase 1: Setup', 'CORS configuration',
     'Configure FastAPI CORS middleware to allow Angular dev server (localhost:4200).',
     'Backend', 'backend/app/main.py',
     'Add CORS middleware to FastAPI main.py allowing origins from .env (default localhost:4200). Allow credentials, all methods, all headers.',
     'To Do'),

    # Phase 2
    (6, 'Phase 2: Models & Auth', 'SQLAlchemy models - Company, User, UserRoleMap',
     'Company model with optional SMTP fields. User model (password varchar 255). UserRoleMap (userId, roleId, companyId, isDefault).',
     'Backend', 'backend/app/models/company.py, backend/app/models/user.py',
     'Create SQLAlchemy models for Company (with nullable MailFrom, MailPassword, SMTP, PortNo), User (userPassword varchar 255 for bcrypt), and UserRoleMap (userId, roleId, companyId, isDefault bit). Add relationships.',
     'To Do'),
    (7, 'Phase 2: Models & Auth', 'SQLAlchemy models - Role, Menu, RoleMenuMap',
     'Role model with IsSuperAdmin. MenuMaster with self-referencing parentMenuId for unlimited nesting. RoleMenuMap with CanAdd/CanRead/CanEdit/CanDelete.',
     'Backend', 'backend/app/models/role.py, backend/app/models/menu.py, backend/app/models/role_menu_map.py',
     'Create SQLAlchemy models for Role (with IsSuperAdmin bit), MenuMaster (self-referencing FK parentMenuId for tree structure, menuOrder for sorting), RoleMenuMap (roleId, menuId, CanAdd, CanRead, CanEdit, CanDelete). Add parent-children relationship on MenuMaster.',
     'To Do'),
    (8, 'Phase 2: Models & Auth', 'SQLAlchemy models - Item, Delivery, Classification, ContactType',
     'ItemGrade, ItemName, ItemLength, ItemSize, DeliveryTerm, DeliveryMode, CustomerClassification, ContactType, CostPointMaster.',
     'Backend', 'backend/app/models/item.py, backend/app/models/delivery.py, backend/app/models/customer_classification.py, backend/app/models/contact_type.py, backend/app/models/cost_point.py',
     'Create SQLAlchemy models for: ItemGrade, ItemName (FK to grade), ItemLength (FK to item), ItemSize (FK to item, corrected: itemSize varchar 50), DeliveryTerm, DeliveryMode, CustomerClassification, ContactType, CostPointMaster (isPrimary, isTax). All have companyId, isActive, audit fields.',
     'To Do'),
    (9, 'Phase 2: Models & Auth', 'SQLAlchemy models - Customer, Enquiry, Quotation, Asset',
     'CustomerMaster, CustomerContacts (FK), CustomerSite (FK). CustomerEnquiry + Details + Costing (with versionNo). QuotSummary (with versionNo, parentQuotId) + Details + T&C. TermsNConditionMaster, RawMaterialCost. Asset.',
     'Backend', 'backend/app/models/customer.py, backend/app/models/enquiry.py, backend/app/models/quotation.py, backend/app/models/terms_condition.py, backend/app/models/raw_material_cost.py, backend/app/models/asset.py',
     'Create SQLAlchemy models for: CustomerMaster, CustomerContacts (FK customerId), CustomerSite (FK customerId). CustomerEnquiry, CustomerEnquiryDetails (FK enqid), CustomerEnquiryCosting (FK enqid + enqdtlid, add versionNo int default 1). QuotSummary (add versionNo, parentQuotId nullable self-FK), QuotDetails (FK quotId), QuotTermsNConditions (FK quotId). TermsNConditionMaster. RawMaterialCost. Asset (FK enqid, quotId).',
     'To Do'),
    (10, 'Phase 2: Models & Auth', 'Alembic initial migration',
     'Generate migration from all 25 models. Run migration to create tables.',
     'Backend', 'backend/alembic/versions/',
     'Generate Alembic migration: alembic revision --autogenerate -m "initial_schema". Review and run: alembic upgrade head. Verify all 25 tables created in SQL Server.',
     'To Do'),
    (11, 'Phase 2: Models & Auth', 'JWT auth - security module',
     'Password hashing with bcrypt. JWT token creation (access + refresh). Token verification. Token payload: { user_id, company_id, role_id, is_super_admin }.',
     'Backend', 'backend/app/core/security.py',
     'Create security.py with: bcrypt password hashing (hash_password, verify_password), JWT token creation (create_access_token, create_refresh_token) with payload { user_id, company_id, role_id, is_super_admin, exp }. Token verification. Use python-jose and passlib[bcrypt].',
     'To Do'),
    (12, 'Phase 2: Models & Auth', 'Auth API endpoints',
     'POST /login (validate credentials, return companies list + temp token). POST /select-company (issue full JWT). POST /switch-company (re-issue JWT). POST /refresh-token.',
     'Backend', 'backend/app/api/v1/auth.py, backend/app/services/auth_service.py',
     'Create auth endpoints: POST /api/v1/auth/login (validate user credentials, return list of user companies from UserRoleMap + temporary token), POST /api/v1/auth/select-company (given temp token + companyId, issue full JWT), POST /api/v1/auth/switch-company (given current JWT + new companyId, re-issue JWT), POST /api/v1/auth/refresh (refresh token flow).',
     'To Do'),
    (13, 'Phase 2: Models & Auth', 'Auth dependencies & middleware',
     'get_current_user dependency (extract from JWT). get_active_company. Permission checker dependency (check RoleMenuMap for menu + action).',
     'Backend', 'backend/app/core/dependencies.py',
     'Create FastAPI dependencies: get_current_user (decode JWT, return user with active company_id and role_id), get_active_company (extract companyId from JWT for query filtering), require_permission(menu_name, action) dependency factory that checks RoleMenuMap for the user role and returns 403 if not allowed.',
     'To Do'),
    (14, 'Phase 2: Models & Auth', 'Frontend - Login page',
     'Login form (username, password). On success: if single company auto-proceed, if multiple show company picker dropdown. Store JWT tokens. Redirect to dashboard.',
     'Frontend', 'frontend/src/app/features/auth/login/',
     'Create Angular login component with reactive form (username, password). Call POST /api/v1/auth/login. If response has multiple companies, show a dropdown to select company, then call POST /select-company. Store access and refresh tokens in localStorage. Redirect to /dashboard.',
     'To Do'),
    (15, 'Phase 2: Models & Auth', 'Frontend - Auth service & interceptor',
     'AuthService (login, selectCompany, switchCompany, logout, isAuthenticated). TokenService (store/get/remove tokens). HTTP interceptor (attach Bearer token, handle 401 -> redirect to login). AuthGuard.',
     'Frontend', 'frontend/src/app/core/auth/',
     'Create AuthService (login, selectCompany, switchCompany, logout, isAuthenticated, getCurrentUser), TokenService (getAccessToken, getRefreshToken, setTokens, clearTokens). HTTP interceptor to add Authorization header, catch 401 -> logout and redirect. AuthGuard (canActivate checks isAuthenticated).',
     'To Do'),

    # Phase 3
    (16, 'Phase 3: Company/User/Role', 'Company CRUD - Backend API',
     'CRUD endpoints for Company. Super Admin only. Includes optional SMTP/email config fields.',
     'Backend', 'backend/app/api/v1/company.py, backend/app/schemas/company.py',
     'Create Company Pydantic schemas (Create, Update, Response) with all fields including nullable SMTP fields. Create CRUD endpoints: GET /api/v1/companies, GET /{id}, POST, PUT /{id}, DELETE /{id} (soft delete). Protect with Super Admin check.',
     'To Do'),
    (17, 'Phase 3: Company/User/Role', 'Company CRUD - Frontend',
     'Company list page with mat-table. Add/Edit form dialog with all fields including optional SMTP configuration section.',
     'Frontend', 'frontend/src/app/features/company/',
     'Create company-list component with Angular Material table (search, pagination). Add/Edit company dialog with reactive form. Include SMTP settings section (MailFrom, MailPassword, SMTP host, PortNo) marked as optional. Only visible to Super Admin.',
     'To Do'),
    (18, 'Phase 3: Company/User/Role', 'Role CRUD - Backend + Frontend',
     'Role CRUD endpoints (company-scoped). Angular list + add/edit dialog.',
     'Full Stack', 'backend/app/api/v1/roles.py, frontend/src/app/features/roles/',
     'Backend: Create Role schemas and CRUD endpoints filtered by companyId. Frontend: Role list page (mat-table) + add/edit dialog (roleName, isSuperAdmin checkbox). Company-scoped.',
     'To Do'),
    (19, 'Phase 3: Company/User/Role', 'Menu Master CRUD - Backend',
     'CRUD endpoints for MenuMaster. GET /menus/tree returns recursive tree JSON. parentMenuId for nesting.',
     'Backend', 'backend/app/api/v1/menus.py, backend/app/services/menu_service.py',
     'Create MenuMaster CRUD endpoints. Key endpoint: GET /api/v1/menus/tree that builds recursive tree from flat menu records (group by parentMenuId, sort by menuOrder). Return nested JSON: [{menuId, menuName, menuOrder, children: [...]}]. Service: build_menu_tree(company_id) method.',
     'To Do'),
    (20, 'Phase 3: Company/User/Role', 'Menu Master CRUD - Frontend',
     'Menu management page. Display tree using mat-tree. Add/edit menu items. Drag-drop reorder. Set parent menu.',
     'Frontend', 'frontend/src/app/features/masters/menu/',
     'Create menu management component using Angular Material mat-tree to display nested menus. Add/edit dialog (menuName, parentMenuId dropdown, menuOrder). Support drag-drop reorder to change menuOrder. Show tree hierarchy.',
     'To Do'),
    (21, 'Phase 3: Company/User/Role', 'Role-Menu Mapping - Backend',
     'GET /role-menu-map/{roleId} returns current permissions. POST /role-menu-map/{roleId} saves permissions as flat array.',
     'Backend', 'backend/app/api/v1/roles.py (or separate role_menu_map.py)',
     'Create endpoints: GET /api/v1/role-menu-map/{roleId} returns list of {menuId, menuName, canAdd, canRead, canEdit, canDelete} for all menus (with current permissions for this role). POST /api/v1/role-menu-map/{roleId} accepts array of {menuId, canAdd, canRead, canEdit, canDelete} and bulk upserts.',
     'To Do'),
    (22, 'Phase 3: Company/User/Role', 'Role-Menu Mapping - Frontend (Tree with Checkboxes)',
     'Tree view using mat-tree. Each node shows menu name + 4 checkboxes (CanAdd, CanRead, CanEdit, CanDelete). Parent checkbox toggles all children. Select role dropdown at top.',
     'Frontend', 'frontend/src/app/features/roles/role-menu-mapping/',
     'Create role-menu-mapping component. Top: role dropdown. Below: mat-tree displaying menu hierarchy. Each tree node has 4 checkboxes (CanAdd, CanRead, CanEdit, CanDelete). Checking a parent toggles all descendants. Unchecking a child unchecks the parent if all children are unchecked. Save button sends flat array to backend.',
     'To Do'),
    (23, 'Phase 3: Company/User/Role', 'User CRUD - Backend',
     'User CRUD endpoints. Include UserRoleMap management (assign user to companies with roles).',
     'Backend', 'backend/app/api/v1/users.py, backend/app/services/user_service.py',
     'Create User CRUD endpoints. On create: hash password with bcrypt. Include sub-resource: POST /users/{id}/role-mappings to assign user to companies with roles (creates UserRoleMap entries). GET /users/{id}/role-mappings returns list of {companyId, companyName, roleId, roleName, isDefault}.',
     'To Do'),
    (24, 'Phase 3: Company/User/Role', 'User CRUD - Frontend',
     'User list page. User form with basic info + company-role mapping section (add multiple company+role pairs).',
     'Frontend', 'frontend/src/app/features/users/',
     'Create user-list component (mat-table with search). User form: basic fields (name, code, email, phone, login, password, reportTo dropdown). Below: company-role mapping section - add rows with company dropdown + role dropdown + isDefault checkbox. Save creates user + role mappings.',
     'To Do'),
    (25, 'Phase 3: Company/User/Role', 'Dynamic sidebar menu',
     'Fetch user-permitted menu tree (only nodes with CanRead=true for user role). Render as recursive mat-tree sidebar. Re-fetch on company switch.',
     'Frontend', 'frontend/src/app/shared/components/dynamic-menu/, frontend/src/app/core/services/menu.service.ts',
     'Create MenuService that calls GET /api/v1/menus/user-tree (backend filters by role permissions, returns only CanRead=true nodes as tree). Create dynamic-menu component using recursive mat-tree in sidenav. Each leaf node links to a route. Re-fetch menu tree when company changes.',
     'To Do'),
    (26, 'Phase 3: Company/User/Role', 'Company Switcher',
     'Dropdown in header showing user companies. On switch: call switch-company API, update JWT, reload menu + current page data.',
     'Frontend', 'frontend/src/app/shared/components/company-switcher/, frontend/src/app/core/services/company-context.service.ts',
     'Create CompanyContextService that tracks active company, provides switchCompany() method. Create company-switcher component (mat-select in toolbar). On switch: call POST /auth/switch-company, store new JWT, emit companyChanged event. All data-fetching components subscribe to companyChanged and reload.',
     'To Do'),
    (27, 'Phase 3: Company/User/Role', 'Permission directive',
     '*hasPermission structural directive that shows/hides elements based on user role permissions for a specific menu + action.',
     'Frontend', 'frontend/src/app/shared/directives/has-permission.directive.ts',
     'Create *hasPermission structural directive. Usage: *hasPermission="\'ItemGrade:CanEdit\'". Directive checks the cached RoleMenuMap permissions (from MenuService) for the current role. If permission is false, remove element from DOM. Also create PermissionGuard for route-level protection.',
     'To Do'),

    # Phase 4
    (28, 'Phase 4: Masters', 'Generic master CRUD pattern - Backend',
     'Create reusable pattern for simple master tables: schema, CRUD endpoints, company-scoped, soft delete, permission-checked.',
     'Backend', 'backend/app/api/v1/masters.py',
     'Create a generic CRUD pattern for master tables. Each master needs: Pydantic schemas (Create, Update, Response), router with GET list (filtered by companyId, isActive), GET by id, POST, PUT, soft DELETE. Apply require_permission dependency. Start with ItemGrade as the template, then replicate for all 11 masters.',
     'To Do'),
    (29, 'Phase 4: Masters', 'Generic master CRUD pattern - Frontend',
     'Create reusable pattern for simple masters: list page (mat-table + search + pagination) + add/edit dialog + delete confirmation.',
     'Frontend', 'frontend/src/app/features/masters/',
     'Create reusable data-table component (mat-table with search, sort, pagination, action buttons). Create confirm-dialog component. For each master: list component using data-table + add/edit dialog with reactive form. Start with item-grade as template. Fields vary per master but pattern is same.',
     'To Do'),
    (30, 'Phase 4: Masters', 'Item Grade CRUD',
     'Backend API + Frontend UI for ItemGrade master.',
     'Full Stack', 'backend/app/api/v1/masters.py, frontend/src/app/features/masters/item-grade/',
     'Implement ItemGrade: Backend schema + CRUD endpoints at /api/v1/masters/item-grades. Frontend: item-grade-list (mat-table) + item-grade-dialog (itemGradeName field). This serves as the template for all other masters.',
     'To Do'),
    (31, 'Phase 4: Masters', 'Item Name CRUD',
     'ItemName with fields: itemGradeName, itemName, itemDia, itemLength, erpItemCode, erpName. Linked to ItemGrade.',
     'Full Stack', 'frontend/src/app/features/masters/item-name/',
     'Implement ItemName CRUD. Form has: itemGradeName (dropdown from ItemGrade), itemName, itemDia, itemLength, erpItemCode, erpName. Backend: /api/v1/masters/item-names.',
     'To Do'),
    (32, 'Phase 4: Masters', 'Item Length & Item Size CRUD',
     'ItemLength (FK itemId, itemLength). ItemSize (FK itemId, itemSize). Both linked to ItemName.',
     'Full Stack', 'frontend/src/app/features/masters/item-length/, frontend/src/app/features/masters/item-size/',
     'Implement ItemLength CRUD (itemId dropdown from ItemName, itemLength field) and ItemSize CRUD (itemId dropdown, itemSize field). Backend endpoints at /api/v1/masters/item-lengths and /api/v1/masters/item-sizes.',
     'To Do'),
    (33, 'Phase 4: Masters', 'Delivery Term & Delivery Mode CRUD',
     'Simple single-field masters: deliveryTerm, deliveryMode.',
     'Full Stack', 'frontend/src/app/features/masters/delivery-term/, frontend/src/app/features/masters/delivery-mode/',
     'Implement DeliveryTerm CRUD (deliveryTerm field) and DeliveryMode CRUD (deliveryMode field). Simple single-field master pattern. Backend: /api/v1/masters/delivery-terms and /delivery-modes.',
     'To Do'),
    (34, 'Phase 4: Masters', 'Contact Type & Customer Classification CRUD',
     'ContactType (contactType). CustomerClassification (classificationName).',
     'Full Stack', 'frontend/src/app/features/masters/contact-type/, frontend/src/app/features/masters/customer-classification/',
     'Implement ContactType CRUD (contactType field) and CustomerClassification CRUD (classificationName field). Simple single-field masters.',
     'To Do'),
    (35, 'Phase 4: Masters', 'Cost Point Master CRUD',
     'CostPointMaster with costPointName, isPrimary (bit), isTax (bit).',
     'Full Stack', 'frontend/src/app/features/masters/cost-point/',
     'Implement CostPointMaster CRUD. Form: costPointName, isPrimary checkbox, isTax checkbox. Backend: /api/v1/masters/cost-points.',
     'To Do'),
    (36, 'Phase 4: Masters', 'Terms & Conditions Master CRUD',
     'TermsNConditionMaster with tncName, tncDescription (varchar 500).',
     'Full Stack', 'frontend/src/app/features/masters/terms-condition/',
     'Implement TermsNConditionMaster CRUD. Form: tncName, tncDescription (textarea). Backend: /api/v1/masters/terms-conditions.',
     'To Do'),
    (37, 'Phase 4: Masters', 'Raw Material Cost CRUD',
     'RawMaterialCost with dia, tpcost (money), effectedFrom (datetime).',
     'Full Stack', 'frontend/src/app/features/masters/raw-material-cost/',
     'Implement RawMaterialCost CRUD. Form: dia (text), tpcost (currency input), effectedFrom (date picker). Backend: /api/v1/masters/raw-material-costs. No lastupdateon/lastupdateby on this table.',
     'To Do'),

    # Phase 5
    (38, 'Phase 5: Customers', 'Customer Master - Backend API',
     'CRUD endpoints for CustomerMaster. Include related contacts and sites in GET response. Search/filter by name, code, classification.',
     'Backend', 'backend/app/api/v1/customers.py, backend/app/schemas/customer.py',
     'Create CustomerMaster schemas and CRUD endpoints. GET /customers supports search by customerName, customerCode, classificationName. GET /customers/{id} includes nested contacts and sites. POST/PUT handle customer basic info only. Separate sub-resource endpoints for contacts and sites.',
     'To Do'),
    (39, 'Phase 5: Customers', 'Customer Contacts - Backend API',
     'CRUD endpoints for CustomerContacts as sub-resource of customer. Multiple contacts per customer.',
     'Backend', 'backend/app/api/v1/customers.py',
     'Create CustomerContacts endpoints: GET /customers/{customerId}/contacts, POST /customers/{customerId}/contacts, PUT /customers/{customerId}/contacts/{contactId}, DELETE (soft). Schema includes contactType, contactPersonName, designation, phones, emails, address, state, dist, birthday, anniversary, office contacts.',
     'To Do'),
    (40, 'Phase 5: Customers', 'Customer Sites - Backend API',
     'CRUD endpoints for CustomerSite as sub-resource. Multiple delivery sites per customer with up to 3 contact persons each.',
     'Backend', 'backend/app/api/v1/customers.py',
     'Create CustomerSite endpoints: GET /customers/{customerId}/sites, POST, PUT, soft DELETE. Schema includes siteAddressCode, addressLine, state, dist, PIN, and 3 sets of contactPerson/Phone/Email fields.',
     'To Do'),
    (41, 'Phase 5: Customers', 'Customer Management - Frontend',
     'Customer list page (mat-table, search, filter by classification). Customer form with 3 tabs: Basic Info, Contacts, Sites. Each tab manages its own sub-table.',
     'Frontend', 'frontend/src/app/features/customers/',
     'Create customer-list (mat-table, search, classification filter). Create customer-form with 3 mat-tabs: Tab 1 (Basic Info: classificationName dropdown, customerCode, customerName, GSTN, PAN, siteId). Tab 2 (Contacts: mat-table with add/edit/delete inline or dialog). Tab 3 (Sites: mat-table with add/edit/delete). Save each tab independently.',
     'To Do'),

    # Phase 6
    (42, 'Phase 6: Enquiries', 'Enquiry CRUD - Backend API',
     'CRUD endpoints for CustomerEnquiry. On customer select, return their contacts and sites. Include nested details.',
     'Backend', 'backend/app/api/v1/enquiries.py, backend/app/services/enquiry_service.py',
     'Create CustomerEnquiry CRUD endpoints. GET /enquiries (list with filters: customer, date range). GET /enquiries/{id} includes nested EnquiryDetails and latest EnquiryCosting. POST creates header only. Helper: GET /customers/{id}/summary returns contacts + sites for dropdown population.',
     'To Do'),
    (43, 'Phase 6: Enquiries', 'Enquiry Details CRUD - Backend API',
     'CRUD for CustomerEnquiryDetails as sub-resource of enquiry. Line items: item, grade, dia, length, unit.',
     'Backend', 'backend/app/api/v1/enquiries.py',
     'Create EnquiryDetails endpoints: GET /enquiries/{enqId}/details, POST (add line item), PUT /{dtlId}, soft DELETE. Each detail has: itemid (FK), itemGradeName, itemDia, itemLength, itemUnit.',
     'To Do'),
    (44, 'Phase 6: Enquiries', 'Enquiry Costing CRUD - Backend API (Versioned)',
     'CRUD for CustomerEnquiryCosting. 20+ cost heads per line item. Version control: create new version duplicates with versionNo+1. Auto-fill TP from RawMaterialCost.',
     'Backend', 'backend/app/api/v1/enquiries.py, backend/app/services/costing_service.py',
     'Create EnquiryCosting endpoints. GET /enquiries/{enqId}/costing?version=latest returns costing for all line items. POST /enquiries/{enqId}/costing creates/updates costing per detail. POST /enquiries/{enqId}/costing/new-version duplicates latest version with versionNo+1. GET /enquiries/{enqId}/costing/versions lists all versions. Auto-fill TPWGST from RawMaterialCost by matching dia. Calculate basicRate, GST (18%), EXFORPrice.',
     'To Do'),
    (45, 'Phase 6: Enquiries', 'Enquiry Management - Frontend',
     'Enquiry list. Enquiry form (customer selector, auto-populate contacts/sites). Enquiry details (editable line items table). Enquiry costing (cost breakdown grid with version controls).',
     'Frontend', 'frontend/src/app/features/enquiries/',
     'Create enquiry-list (mat-table, filters). Enquiry-form: customer dropdown (on select, populate contact and site dropdowns), enqDate picker, enqMode, description, validityDays. Enquiry-details: editable mat-table for line items (item dropdown, grade, dia, length, unit, add/remove rows). Enquiry-costing: grid showing all 20+ cost heads per line item, auto-calculate totals. Version dropdown to switch versions. "Create New Version" button. Previous versions read-only.',
     'To Do'),

    # Phase 7
    (46, 'Phase 7: Quotations', 'Quotation CRUD - Backend API (Versioned)',
     'CRUD for QuotSummary with versioning. Link to enquiry, customer, site, contact. Revision creates new row with versionNo+1, parentQuotId.',
     'Backend', 'backend/app/api/v1/quotations.py, backend/app/services/quotation_service.py, backend/app/services/versioning_service.py',
     'Create QuotSummary CRUD endpoints. GET /quotations (list, show latest version per parentQuotId group). GET /{id} with nested details + T&C. POST creates quotation linked to enquiry. POST /quotations/{id}/revise creates new row with versionNo+1, parentQuotId=original, quotNo format QUOT-XXX-R{n}. GET /quotations/{id}/versions lists all versions. Approval: PUT /quotations/{id}/approve sets approvedby, approvedon (only latest version).',
     'To Do'),
    (47, 'Phase 7: Quotations', 'Quotation Details CRUD - Backend',
     'Line items for quotation: grade, dia, length, unit, quantity, basicRate, IGST, CGST, SGST, totAmount, totRate.',
     'Backend', 'backend/app/api/v1/quotations.py',
     'Create QuotDetails endpoints as sub-resource of quotation. GET /quotations/{quotId}/details, POST, PUT, soft DELETE. Fields: itemGradeName, itemDia, itemLength, itemUnit, quantity, basicRate, IGST, CGST, SGST, totAmount (calculated), totRate (calculated).',
     'To Do'),
    (48, 'Phase 7: Quotations', 'Quotation T&C - Backend',
     'Terms & Conditions per quotation. Pull defaults from TermsNConditionMaster, allow per-quotation edit.',
     'Backend', 'backend/app/api/v1/quotations.py',
     'Create QuotTermsNConditions endpoints. GET /quotations/{quotId}/terms. POST /quotations/{quotId}/terms/from-master copies all active T&C from TermsNConditionMaster to this quotation. PUT /{tncId} allows editing per-quotation. Soft DELETE.',
     'To Do'),
    (49, 'Phase 7: Quotations', 'Quotation Management - Frontend',
     'Quotation list (latest versions). Quotation form with customer/enquiry linking. Details table. T&C editor. Version history panel.',
     'Frontend', 'frontend/src/app/features/quotations/',
     'Create quotation-list (mat-table showing latest version per quotation, filters). Quotation-form: link to enquiry (dropdown, auto-fills customer/site/contact), quotDate, subject, deliveryTerm/Mode dropdowns, refQuotNo, remarks, CustomerPONo/Date. Quotation-details: editable table (grade, dia, length, qty, rates, GST split, totals). Quotation-tnc: list with add from master button + inline edit. Version-history: side panel listing all versions, click to view (read-only for old versions). "Revise Quotation" button.',
     'To Do'),
    (50, 'Phase 7: Quotations', 'Quotation Print/PDF',
     'Formatted printable quotation view for a specific version. Include header, details table, T&C, footer.',
     'Frontend', 'frontend/src/app/features/quotations/quotation-print/',
     'Create quotation-print component that renders a formatted quotation for a specific version. Include: company header, quotation details (quotNo, date, customer, site, contact), line items table with totals, T&C section. Use @media print CSS. Add "Print" button that triggers window.print(). Optionally generate PDF using jsPDF or similar.',
     'To Do'),

    # Phase 8
    (51, 'Phase 8: Assets & Email', 'Azure Blob Service',
     'Service for uploading, downloading, and deleting files from Azure Blob Storage. Connection from .env.',
     'Backend', 'backend/app/services/azure_blob_service.py',
     'Create AzureBlobService using azure-storage-blob SDK. Methods: upload_file(file, container, blob_name), download_file(blob_name), delete_file(blob_name), generate_sas_url(blob_name). Connection string from .env AZURE_BLOB_CONNECTION_STRING. Container from AZURE_BLOB_CONTAINER.',
     'To Do'),
    (52, 'Phase 8: Assets & Email', 'Asset CRUD - Backend + Frontend',
     'Upload files linked to enquiry or quotation. List assets. Download. Soft delete.',
     'Full Stack', 'backend/app/api/v1/assets.py, frontend/src/app/features/assets/',
     'Backend: POST /api/v1/assets/upload (multipart file + enqid/quotId). GET /assets?enqid=X or ?quotId=X. GET /assets/{id}/download (returns SAS URL or streams file). Soft DELETE. Frontend: asset-upload component (drag-drop or file picker), asset-list component (table with download/delete actions). Integrate into enquiry and quotation detail pages.',
     'To Do'),
    (53, 'Phase 8: Assets & Email', 'Email Configuration - Super Admin UI',
     'Company edit form: SMTP settings section (host, port, MailFrom, MailPassword). All optional. If not set, email features hidden.',
     'Frontend', 'frontend/src/app/features/company/ (update existing)',
     'Update company edit form to include an "Email Configuration" section: SMTP Host, Port, MailFrom (From email address), MailPassword (password input). All fields optional. Add a "Test Connection" button that calls backend to verify SMTP settings. If SMTP not configured for a company, hide all email-related buttons across the app.',
     'To Do'),
    (54, 'Phase 8: Assets & Email', 'Mail Template Management',
     'Super Admin can create/edit HTML email templates per company. Templates used when sending quotation emails.',
     'Full Stack', 'backend/app/api/v1/email.py, frontend/src/app/features/company/ (or dedicated)',
     'Create a simple MailTemplate model (or use a config table) storing template name + HTML body per company. Backend CRUD for templates. Frontend: template editor (textarea or basic rich text) in Super Admin panel. Default template provided. Template supports placeholders like {{customerName}}, {{quotNo}}, {{quotDate}}.',
     'To Do'),
    (55, 'Phase 8: Assets & Email', 'Send Quotation Email - Backend + Frontend',
     'Send quotation PDF to customer contact email. From = company MailFrom. To = selected contact email. Uses company SMTP config.',
     'Full Stack', 'backend/app/services/email_service.py, backend/app/api/v1/email.py',
     'Backend: EmailService that loads company SMTP config, constructs email (From=company MailFrom, To=contact email, subject from template, body from template with placeholders filled, attach quotation PDF). POST /api/v1/email/send-quotation {quotId, contactId, templateId}. Frontend: "Send Email" button on quotation page, shows dialog to confirm recipient (pre-filled from selected contact), select template, send. Button hidden if company SMTP not configured.',
     'To Do'),
]

# Write rows
for row_num, task in enumerate(tasks, 2):
    for col_num, value in enumerate(task, 1):
        cell = ws.cell(row=row_num, column=col_num, value=value)
        cell.font = normal_font
        cell.alignment = wrap_alignment
        cell.border = thin_border
        if col_num == 8:  # Status column
            cell.alignment = center_alignment

# Add conditional formatting / colors for phases
phase_colors = {
    'Phase 1: Setup': 'E8F0FE',
    'Phase 2: Models & Auth': 'FFF2CC',
    'Phase 3: Company/User/Role': 'D9EAD3',
    'Phase 4: Masters': 'FCE5CD',
    'Phase 5: Customers': 'D9D2E9',
    'Phase 6: Enquiries': 'CFE2F3',
    'Phase 7: Quotations': 'F4CCCC',
    'Phase 8: Assets & Email': 'EAD1DC',
}

for row_num in range(2, len(tasks) + 2):
    phase = ws.cell(row=row_num, column=2).value
    if phase in phase_colors:
        fill = PatternFill(start_color=phase_colors[phase], end_color=phase_colors[phase], fill_type='solid')
        for col_num in range(1, 9):
            ws.cell(row=row_num, column=col_num).fill = fill

# Set row heights
for row_num in range(2, len(tasks) + 2):
    ws.row_dimensions[row_num].height = 60

# Auto-filter
ws.auto_filter.ref = f"A1:H{len(tasks) + 1}"

wb.save('tentative-implementation-plan.xlsx')
print("Created: tentative-implementation-plan.xlsx")
print(f"\nTotal tasks: {len(tasks)}")
